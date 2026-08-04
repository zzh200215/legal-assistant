from __future__ import annotations

import mimetypes
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.models.document import Document
from app.models.user import User
from app.services.storage_service import storage_service


class DocumentDeliveryError(ValueError):
    pass


class DocumentDeliveryService:
    """Creates short-lived, attributed copies for governed document downloads."""

    TEXT_FORMATS = {"txt", "md", "csv"}
    WATERMARKABLE_FORMATS = TEXT_FORMATS | {"docx", "xlsx"}

    def _source_path(self, document: Document) -> Path:
        root = storage_service.base_dir().resolve()
        source = Path(document.file_path).resolve()
        try:
            source.relative_to(root)
        except ValueError as exc:
            raise DocumentDeliveryError("文档源文件不在受控存储目录中") from exc
        if not source.is_file():
            raise DocumentDeliveryError("文档源文件不存在或已被清理")
        return source

    @staticmethod
    def _safe_filename(document: Document, source: Path) -> str:
        suffix = source.suffix or f".{document.file_type.lstrip('.')}"
        title = Path(document.title or "document").name.replace("/", "_").replace("\\", "_")
        return title if title.lower().endswith(suffix.lower()) else f"{title}{suffix}"

    @staticmethod
    def _watermark_text(document: Document, user: User) -> str:
        identity = user.full_name or user.username or f"用户{user.id}"
        timestamp = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M")
        return f"内部受控副本 | 文档#{document.id} | 下载人：{identity} | 下载时间：{timestamp}"

    def _temp_path(self, source: Path) -> Path:
        directory = storage_service.ensure_dir(storage_service.base_dir() / "governed_downloads")
        return directory / f"{uuid.uuid4().hex}{source.suffix.lower()}"

    @staticmethod
    def _watermark_text_file(source: Path, target: Path, watermark: str) -> None:
        if source.suffix.lower() == ".csv":
            target.write_text(f"# {watermark}\n" + source.read_text(encoding="utf-8-sig"), encoding="utf-8")
            return
        target.write_text(f"<!-- {watermark} -->\n\n" + source.read_text(encoding="utf-8"), encoding="utf-8")

    @staticmethod
    def _watermark_docx(source: Path, target: Path, watermark: str) -> None:
        doc = DocxDocument(source)
        for section in doc.sections:
            paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
            paragraph.text = watermark
        doc.save(target)

    @staticmethod
    def _watermark_xlsx(source: Path, target: Path, watermark: str) -> None:
        workbook = load_workbook(source)
        for worksheet in workbook.worksheets:
            worksheet.oddHeader.center.text = watermark
        workbook.save(target)

    def prepare_download(self, *, document: Document, user: User) -> dict:
        if not document.download_enabled:
            raise DocumentDeliveryError("该文档已被设置为禁止下载")
        source = self._source_path(document)
        filename = self._safe_filename(document, source)
        extension = source.suffix.lower().lstrip(".")
        result = {
            "path": source,
            "filename": filename,
            "media_type": mimetypes.guess_type(filename)[0] or "application/octet-stream",
            "watermark_applied": False,
            "watermark_supported": True,
            "temporary": False,
        }
        if not document.watermark_required:
            return result
        if extension not in self.WATERMARKABLE_FORMATS:
            result["watermark_supported"] = False
            return result

        target = self._temp_path(source)
        try:
            watermark = self._watermark_text(document, user)
            if extension in self.TEXT_FORMATS:
                self._watermark_text_file(source, target, watermark)
            elif extension == "docx":
                self._watermark_docx(source, target, watermark)
            else:
                self._watermark_xlsx(source, target, watermark)
        except Exception as exc:
            target.unlink(missing_ok=True)
            raise DocumentDeliveryError("生成受控下载副本失败") from exc
        result.update(path=target, watermark_applied=True, temporary=True)
        return result

    @staticmethod
    def cleanup(path: Path) -> None:
        try:
            path.unlink(missing_ok=True)
        except OSError:
            pass


document_delivery_service = DocumentDeliveryService()
