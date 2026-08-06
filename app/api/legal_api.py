import asyncio
import json
import logging
import os
import re
import tempfile
import csv
from pathlib import Path
from datetime import datetime, date

import jieba
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from app.core.api_response import api_error, should_passthrough_exception
from app.core.auth import get_current_user
from app.core.database import get_db
from app.core.config import get_settings
from app.core.time import utc_now
from app.models.legal import ContractReview, LegalArticle, LegalConsultation, LegalDraft, LegalSource
from app.models.user import User
from app.services.legal_service import (
    DISCLAIMER,
    compare_contracts,
    ensure_demo_sources,
    target_query,
    test_retrieval,
)
from app.services.audit_log_service import AuditLogService
from app.services.legal_retrieval_service import legal_retrieval_service
from app.services.legal_knowledge_graph_service import legal_knowledge_graph_service
from app.services.legal_workspace_service import (
    legal_workspace_module,
    legal_workspace_read_module,
    serialize_workspace_row,
)

audit = AuditLogService()
logger = logging.getLogger(__name__)

router = APIRouter()


def _run_async(coro):
    """E-7：在请求线程（sync 端点由 FastAPI 调度到线程池）中运行 async 服务调用。

    async 端点中的同步 DB 调用会阻塞事件循环，拖垮并发；sync 端点由
    FastAPI 放入线程池执行，DB Session 全程同一线程，事件循环不再被阻塞。
    """
    return asyncio.run(coro)

class ConsultationIn(BaseModel):
    question: str = Field(min_length=1, max_length=12000)
    case_id: int | None = None


class ContractReviewIn(BaseModel):
    title: str = Field(default="未命名合同", max_length=256)
    content: str = Field(min_length=1, max_length=50000)
    document_id: int | None = None
    review_policy_id: int | None = None
    # 允许前端基于模板编辑一次性副本，但不允许伪造模板版本。
    review_policy_override: dict | None = None
    case_id: int | None = None


class DraftIn(BaseModel):
    document_type: str
    fields: dict[str, str] = {}
    case_id: int | None = None


class ReviewActionIn(BaseModel):
    action: str
    note: str | None = Field(default=None, max_length=2000)


class ReviewCommentIn(BaseModel):
    note: str = Field(min_length=1, max_length=2000)


def serialize(row):
    return serialize_workspace_row(row)


@router.get("/overview")
def overview(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return legal_workspace_read_module.overview(db, current_user)


@router.get("/metrics")
def legal_metrics(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return legal_workspace_read_module.metrics(db, current_user)


@router.get("/sources")
def list_sources(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    ensure_demo_sources(db, current_user.id)
    return db.query(LegalSource).filter(LegalSource.user_id == current_user.id).order_by(LegalSource.updated_at.desc()).all()


@router.post("/sources/import")
async def import_sources(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """批量导入法源资料（CSV/Excel）。

    CSV 格式要求：
    - 必需列：title, source_type, content
    - 可选列：citation, jurisdiction, version, effective_date, status
    - source_type: statute, case, template
    - status: active, inactive, pending_update
    """
    if current_user.role not in {"admin", "dept_admin"}:
        raise api_error(403, "仅管理员可批量导入法源", code="LEGAL_SOURCE_IMPORT_FORBIDDEN")

    allowed_exts = {".csv", ".xlsx", ".xls"}
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in allowed_exts:
        raise api_error(400, f"不支持的文件类型: {ext}，支持 CSV/Excel", code="LEGAL_SOURCE_FILE_TYPE_INVALID")

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            content_bytes = await file.read()
            tmp.write(content_bytes)
            tmp_path = tmp.name

        rows = []
        if ext == ".csv":
            with open(tmp_path, "r", encoding="utf-8-sig") as f:
                reader = csv.DictReader(f)
                rows = list(reader)
        else:
            try:
                from openpyxl import load_workbook
                wb = load_workbook(tmp_path, read_only=True, data_only=True)
                ws = wb.active
                sheet_rows = list(ws.iter_rows(values_only=True))
                wb.close()
            except ImportError:
                raise api_error(500, "Excel 导入需要安装 openpyxl，请使用 CSV 格式", code="EXCEL_LIBRARY_MISSING")
            if len(sheet_rows) < 2:
                raise api_error(400, "Excel 文件为空或缺少数据行", code="LEGAL_SOURCE_FILE_EMPTY")
            headers = [str(h).strip() if h is not None else "" for h in sheet_rows[0]]
            for data_row in sheet_rows[1:]:
                row_dict = {headers[i]: (data_row[i] if i < len(data_row) else None) for i in range(len(headers))}
                if any(v not in (None, "") for v in row_dict.values()):
                    rows.append(row_dict)

        if not rows:
            raise api_error(400, "文件为空或格式错误", code="LEGAL_SOURCE_FILE_EMPTY")

        # 校验必需列
        required_cols = {"title", "source_type", "content"}
        first_row_keys = set(rows[0].keys())
        missing_cols = required_cols - first_row_keys
        if missing_cols:
            raise api_error(400, f"缺少必需列: {', '.join(missing_cols)}", code="LEGAL_SOURCE_MISSING_COLUMNS")

        # 批量插入
        imported = 0
        skipped = 0
        errors = []

        def _cell(value, default=""):
            if value is None:
                return default
            return str(value).strip()

        for idx, row in enumerate(rows, start=1):
            try:
                title = _cell(row.get("title"))
                source_type = _cell(row.get("source_type"), "statute")
                content = _cell(row.get("content"))

                if not title or not content:
                    skipped += 1
                    errors.append(f"第 {idx} 行：标题或内容为空")
                    continue

                if source_type not in {"statute", "case", "template"}:
                    source_type = "statute"

                # 解析生效日期（支持字符串 YYYY-MM-DD 或 Excel date 对象）
                effective_date_raw = row.get("effective_date")
                effective_date_val = None
                if isinstance(effective_date_raw, date):
                    effective_date_val = effective_date_raw
                elif isinstance(effective_date_raw, datetime):
                    effective_date_val = effective_date_raw.date()
                else:
                    effective_date_str = _cell(effective_date_raw)
                    if effective_date_str:
                        try:
                            effective_date_val = datetime.strptime(effective_date_str, "%Y-%m-%d").date()
                        except ValueError:
                            pass

                status = _cell(row.get("status"), "active")
                if status not in {"active", "inactive", "pending_update"}:
                    status = "active"

                source = LegalSource(
                    user_id=current_user.id,
                    title=title,
                    source_type=source_type,
                    citation=_cell(row.get("citation")) or None,
                    jurisdiction=_cell(row.get("jurisdiction"), "中国大陆"),
                    effective_date=effective_date_val,
                    version=_cell(row.get("version"), "v1"),
                    status=status,
                    content=content,
                )
                db.add(source)
                imported += 1
            except Exception as e:
                skipped += 1
                errors.append(f"第 {idx} 行：{str(e)[:100]}")

        db.commit()
        audit.log(db, current_user, "legal_source_import", detail=f"imported={imported}, skipped={skipped}")

        return {
            "imported": imported,
            "skipped": skipped,
            "total": len(rows),
            "errors": errors[:10],  # 最多返回前 10 条错误
        }

    except HTTPException:
        raise
    except Exception as exc:
        if should_passthrough_exception(exc):
            raise
        raise api_error(500, "法源导入失败", code="LEGAL_SOURCE_IMPORT_FAILED", detail=str(exc))
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


class RetrievalTestIn(BaseModel):
    question: str = Field(min_length=1, max_length=2000)


@router.post("/sources/retrieval-test")
def retrieval_test(req: RetrievalTestIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """检索测试工具：输入问题实时查看法源召回排序及评分明细（FL.md 6.1）。"""
    ensure_demo_sources(db, current_user.id)
    sources = db.query(LegalSource).filter(LegalSource.user_id == current_user.id).all()
    results = test_retrieval(req.question, sources)
    return {"question": req.question, "results": results, "total_sources": len(sources)}


@router.post("/sources/hybrid-retrieval-test")
def hybrid_retrieval_test(req: RetrievalTestIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """条文级混合检索测试工具（词法+向量+RRF融合）——对应 /article-search 生产路径。"""
    ensure_demo_sources(db, current_user.id)
    results = _run_async(legal_retrieval_service.search(db, req.question.strip(), current_user.id, limit=10))
    return {"question": req.question, "results": results, "total_articles": len(results)}


class SourceStatusUpdateIn(BaseModel):
    status: str = Field(pattern="^(active|inactive|pending_update)$")


@router.patch("/sources/{source_id}/status")
async def update_source_status(
    source_id: int,
    req: SourceStatusUpdateIn,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """更新法源状态（版本管理）。"""
    source = db.query(LegalSource).filter(LegalSource.id == source_id, LegalSource.user_id == current_user.id).first()
    if not source:
        raise api_error(404, "法源不存在", code="LEGAL_SOURCE_NOT_FOUND")

    old_status = source.status
    source.status = req.status
    source.updated_at = utc_now()
    db.commit()
    graph_synced = await legal_knowledge_graph_service.sync_source(db, source.id, current_user.id)
    audit.log(db, current_user, "legal_source_status_update", target_type="source", target_id=source_id, detail=f"{old_status} -> {req.status}")

    return {"id": source_id, "status": source.status, "updated_at": source.updated_at, "graph_synced": graph_synced}


class SourceCreateIn(BaseModel):
    title: str = Field(min_length=1, max_length=256)
    source_type: str = Field(pattern="^(statute|case|template|judicial_interpretation)$")
    content: str = Field(min_length=1)
    citation: str | None = Field(default=None, max_length=256)
    jurisdiction: str = Field(default="中国大陆", max_length=128)
    version: str = Field(default="v1", max_length=64)
    status: str = Field(default="active", pattern="^(active|inactive|pending_update)$")
    # V2.0 新增字段
    document_number: str | None = Field(default=None, max_length=64)
    promulgator: str | None = Field(default=None, max_length=128)
    full_text: str | None = Field(default=None)
    law_areas: list[str] = Field(default_factory=list)
    keywords: list[str] = Field(default_factory=list)
    amends: list[int] = Field(default_factory=list, description="被当前法源修订的法源 ID")
    amended_by: list[int] = Field(default_factory=list, description="修订当前法源的法源 ID")


class SourceUpdateIn(BaseModel):
    title: str = Field(min_length=1, max_length=256)
    source_type: str = Field(pattern="^(statute|case|template|judicial_interpretation)$")
    content: str = Field(min_length=1)
    citation: str | None = Field(default=None, max_length=256)
    jurisdiction: str = Field(default="中国大陆", max_length=128)
    version: str = Field(default="v1", max_length=64)
    status: str = Field(pattern="^(active|inactive|pending_update)$")
    document_number: str | None = Field(default=None, max_length=64)
    promulgator: str | None = Field(default=None, max_length=128)
    full_text: str | None = Field(default=None)
    law_areas: list[str] = Field(default_factory=list)
    keywords: list[str] = Field(default_factory=list)
    amends: list[int] = Field(default_factory=list, description="被当前法源修订的法源 ID")
    amended_by: list[int] = Field(default_factory=list, description="修订当前法源的法源 ID")


def _serialize_source(source: LegalSource) -> dict:
    return {
        "id": source.id,
        "title": source.title,
        "source_type": source.source_type,
        "citation": source.citation,
        "jurisdiction": source.jurisdiction,
        "effective_date": source.effective_date,
        "version": source.version,
        "status": source.status,
        "content": source.content,
        "created_at": source.created_at,
        "updated_at": source.updated_at,
        # V2.0 扩展字段
        "document_number": source.document_number,
        "promulgator": source.promulgator,
        "promulgation_date": source.promulgation_date,
        "full_text": source.full_text,
        "law_areas": json.loads(source.law_area_json) if source.law_area_json else [],
        "keywords": json.loads(source.keywords_json) if source.keywords_json else [],
        "amended_by": json.loads(source.amended_by_json) if source.amended_by_json else [],
        "amends": json.loads(source.amends_json) if source.amends_json else [],
    }


@router.post("/sources")
async def create_source(req: SourceCreateIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """单条创建法源（管理后台，替代直接操作数据库）。"""
    if current_user.role not in {"admin", "dept_admin"}:
        raise api_error(403, "仅管理员可创建法源", code="LEGAL_SOURCE_CREATE_FORBIDDEN")
    source = LegalSource(
        user_id=current_user.id, title=req.title, source_type=req.source_type,
        content=req.content, citation=req.citation, jurisdiction=req.jurisdiction,
        version=req.version, status=req.status, effective_date=date.today(),
        document_number=req.document_number, promulgator=req.promulgator,
        full_text=req.full_text,
        law_area_json=json.dumps(req.law_areas, ensure_ascii=False) if req.law_areas else None,
        keywords_json=json.dumps(req.keywords, ensure_ascii=False) if req.keywords else None,
        amends_json=json.dumps(req.amends) if req.amends else None,
        amended_by_json=json.dumps(req.amended_by) if req.amended_by else None,
    )
    db.add(source)
    db.commit()
    db.refresh(source)
    graph_synced = await legal_knowledge_graph_service.sync_source(db, source.id, current_user.id)
    result = _serialize_source(source)
    result["graph_synced"] = graph_synced
    audit.log(db, current_user, "legal_source_create", target_type="source", target_id=source.id, detail=req.title)
    return result


@router.put("/sources/{source_id}")
async def update_source(source_id: int, req: SourceUpdateIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """编辑法源内容（管理后台）。"""
    if current_user.role not in {"admin", "dept_admin"}:
        raise api_error(403, "仅管理员可编辑法源", code="LEGAL_SOURCE_EDIT_FORBIDDEN")
    source = db.query(LegalSource).filter(LegalSource.id == source_id, LegalSource.user_id == current_user.id).first()
    if not source:
        raise api_error(404, "法源不存在", code="LEGAL_SOURCE_NOT_FOUND")

    source.title = req.title
    source.source_type = req.source_type
    source.content = req.content
    source.citation = req.citation
    source.jurisdiction = req.jurisdiction
    source.version = req.version
    source.status = req.status
    source.document_number = req.document_number
    source.promulgator = req.promulgator
    source.full_text = req.full_text
    source.law_area_json = json.dumps(req.law_areas, ensure_ascii=False) if req.law_areas else None
    source.keywords_json = json.dumps(req.keywords, ensure_ascii=False) if req.keywords else None
    source.amends_json = json.dumps(sorted(set(item for item in req.amends if item != source_id))) if req.amends else None
    source.amended_by_json = json.dumps(sorted(set(item for item in req.amended_by if item != source_id))) if req.amended_by else None
    source.updated_at = utc_now()
    db.commit()
    db.refresh(source)
    graph_synced = await legal_knowledge_graph_service.sync_source(db, source.id, current_user.id)
    result = _serialize_source(source)
    result["graph_synced"] = graph_synced
    audit.log(db, current_user, "legal_source_update", target_type="source", target_id=source_id, detail=req.title)
    return result


@router.delete("/sources/{source_id}")
async def delete_source(source_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """删除法源（管理后台）。"""
    if current_user.role not in {"admin", "dept_admin"}:
        raise api_error(403, "仅管理员可删除法源", code="LEGAL_SOURCE_DELETE_FORBIDDEN")
    source = db.query(LegalSource).filter(LegalSource.id == source_id, LegalSource.user_id == current_user.id).first()
    if not source:
        raise api_error(404, "法源不存在", code="LEGAL_SOURCE_NOT_FOUND")

    title = source.title
    db.delete(source)
    db.commit()
    await legal_knowledge_graph_service.delete_source(source_id, current_user.id)
    audit.log(db, current_user, "legal_source_delete", target_type="source", target_id=source_id, detail=title)
    return {"deleted": True, "id": source_id}


# ── 条文级检索 ───────────────────────────────────────────────


@router.get("/sources/{source_id}/articles")
def list_articles(source_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """查看指定法源的条文列表。"""
    source = db.query(LegalSource).filter(LegalSource.id == source_id, LegalSource.user_id == current_user.id).first()
    if not source:
        raise api_error(404, "法源不存在", code="LEGAL_SOURCE_NOT_FOUND")
    articles = (
        db.query(LegalArticle)
        .filter(LegalArticle.source_id == source_id)
        .order_by(LegalArticle.sequence)
        .all()
    )
    return [
        {
            "id": a.id,
            "article_number": a.article_number,
            "title": a.title,
            "content": a.content,
            "chapter": a.chapter,
            "section": a.section,
            "sequence": a.sequence,
        }
        for a in articles
    ]


@router.get("/sources/{source_id}/relations")
def source_relations(
    source_id: int,
    depth: int = 1,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """法条关系图谱：BFS 多跳遍历修订关系链（depth 最大3层）。"""
    if depth < 1 or depth > 3:
        raise api_error(400, "depth 范围为 1-3", code="INVALID_DEPTH")

    root = db.query(LegalSource).filter(
        LegalSource.id == source_id, LegalSource.user_id == current_user.id
    ).first()
    if not root:
        raise api_error(404, "法源不存在", code="LEGAL_SOURCE_NOT_FOUND")

    def _node(s: LegalSource) -> dict:
        return {
            "id": s.id,
            "title": s.title,
            "citation": s.citation,
            "document_number": s.document_number,
            "effective_date": s.effective_date,
            "status": s.status,
            "version": s.version,
        }

    def _direct_relations(raw: str | None) -> list[dict]:
        try:
            relation_ids = json.loads(raw) if raw else []
        except (TypeError, ValueError):
            relation_ids = []
        nodes = []
        for relation_id in relation_ids:
            related = db.query(LegalSource).filter(
                LegalSource.id == relation_id,
                LegalSource.user_id == current_user.id,
            ).first()
            if related:
                nodes.append(_node(related))
        return nodes

    def _bfs(start_id: int, direction: str) -> list[dict]:
        """BFS 多跳遍历。direction='amended_by'|'amends'"""
        edges = []
        visited: set[int] = {start_id}
        queue: list[tuple[int, int]] = [(start_id, 0)]
        while queue:
            cur_id, cur_depth = queue.pop(0)
            if cur_depth >= depth:
                continue
            node = db.query(LegalSource).filter(
                LegalSource.id == cur_id, LegalSource.user_id == current_user.id
            ).first()
            if not node:
                continue
            raw = node.amended_by_json if direction == "amended_by" else node.amends_json
            neighbor_ids: list[int] = json.loads(raw) if raw else []
            for nid in neighbor_ids:
                edge: dict = {"from": cur_id, "to": nid, "depth": cur_depth + 1}
                neighbor = db.query(LegalSource).filter(
                    LegalSource.id == nid, LegalSource.user_id == current_user.id
                ).first()
                if neighbor:
                    edge["node"] = _node(neighbor)
                    if nid not in visited:
                        visited.add(nid)
                        queue.append((nid, cur_depth + 1))
                else:
                    edge["node"] = None
                edges.append(edge)
        return edges

    return {
        "source_id": source_id,
        "title": root.title,
        "depth": depth,
        "amended_by": _direct_relations(root.amended_by_json),
        "amends": _direct_relations(root.amends_json),
        "amended_by_edges": _bfs(source_id, "amended_by"),
        "amends_edges": _bfs(source_id, "amends"),
    }


@router.get("/sources/{source_id}/knowledge-graph")
async def source_knowledge_graph(
    source_id: int,
    depth: int = 2,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Return the Neo4j version-relation graph for one legal source."""
    if depth < 1 or depth > 3:
        raise api_error(400, "depth 范围为 1-3", code="INVALID_DEPTH")
    source = db.query(LegalSource).filter(
        LegalSource.id == source_id, LegalSource.user_id == current_user.id
    ).first()
    if not source:
        raise api_error(404, "法源不存在", code="LEGAL_SOURCE_NOT_FOUND")
    graph = await legal_knowledge_graph_service.source_graph(current_user.id, source_id, depth)
    return {"source_id": source_id, "depth": depth, **graph}


@router.get("/knowledge-graph/health")
async def knowledge_graph_health(current_user: User = Depends(get_current_user)):
    """Report graph availability without exposing Neo4j credentials."""
    if current_user.role not in {"admin", "dept_admin"}:
        raise api_error(403, "仅管理员可查看图谱状态", code="LEGAL_GRAPH_HEALTH_FORBIDDEN")
    return await legal_knowledge_graph_service.health()


@router.post("/knowledge-graph/reindex")
async def reindex_knowledge_graph(
    source_id: int | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Backfill Neo4j from the relational legal-source data."""
    if current_user.role not in {"admin", "dept_admin"}:
        raise api_error(403, "仅管理员可重建法律知识图谱", code="LEGAL_GRAPH_REINDEX_FORBIDDEN")
    if source_id is not None:
        source = db.query(LegalSource).filter(
            LegalSource.id == source_id, LegalSource.user_id == current_user.id
        ).first()
        if not source:
            raise api_error(404, "法源不存在", code="LEGAL_SOURCE_NOT_FOUND")
        source_count = 1 if await legal_knowledge_graph_service.sync_source(db, source_id, current_user.id) else 0
    else:
        source_count = await legal_knowledge_graph_service.sync_sources(db, current_user.id)
    audit.log(db, current_user, "legal_knowledge_graph_reindex", detail=f"sources={source_count}")
    return {"synced_sources": source_count, "graph_enabled": legal_knowledge_graph_service.enabled}


@router.get("/article-search")
def search_articles(q: str = "", db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """条文级混合检索：精确/关键词召回与向量语义召回融合。"""
    if not q or len(q.strip()) < 2:
        raise api_error(400, "搜索关键词至少2个字符", code="ARTICLE_SEARCH_TOO_SHORT")
    return _run_async(legal_retrieval_service.search(db, q.strip(), current_user.id))


@router.post("/sources/{source_id}/reindex")
async def reindex_source_articles(source_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """写入经审核法源的条文向量；不影响已有的文档 RAG 索引。"""
    if current_user.role not in {"admin", "dept_admin"}:
        raise api_error(403, "仅管理员可重建法源索引", code="LEGAL_SOURCE_REINDEX_FORBIDDEN")
    source = db.query(LegalSource).filter(LegalSource.id == source_id, LegalSource.user_id == current_user.id).first()
    if not source:
        raise api_error(404, "法源不存在", code="LEGAL_SOURCE_NOT_FOUND")
    try:
        indexed = await legal_retrieval_service.index_source(db, source_id, current_user.id)
    except Exception as exc:
        logger.warning("Legal source vector index failed for source %s: %s", source_id, type(exc).__name__)
        raise api_error(503, "法源向量索引暂不可用，请检查 Embedding 服务后重试", code="LEGAL_SOURCE_REINDEX_UNAVAILABLE")
    graph_synced = await legal_knowledge_graph_service.sync_source(db, source_id, current_user.id)
    audit.log(
        db,
        current_user,
        "legal_source_reindex",
        target_type="source",
        target_id=source_id,
        detail=f"articles={indexed}; graph_synced={graph_synced}",
    )
    return {"source_id": source_id, "indexed_articles": indexed, "graph_synced": graph_synced}


@router.post("/consultations")
def create_consultation(req: ConsultationIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        operation = _run_async(
            legal_workspace_module.create_consultation(db, current_user, req.question, case_id=req.case_id)
        )
    except ValueError as exc:
        if str(exc) == "QUOTA_EXCEEDED":
            raise api_error(429, "本月咨询配额已用完，请升级订阅", code="QUOTA_EXCEEDED")
        raise
    except LookupError as exc:
        if str(exc) == "LEGAL_CASE_NOT_FOUND":
            raise api_error(404, "案件不存在或无权访问", code="LEGAL_CASE_NOT_FOUND")
        raise
    result = serialize(operation.row)
    result["disclaimer_level"] = operation.disclaimer["level"]
    result["disclaimer_label"] = operation.disclaimer["label"]
    return result


class FollowupIn(BaseModel):
    question: str = Field(min_length=1, max_length=5000)


@router.post("/consultations/{item_id}/followup")
def followup_consultation(item_id: int, req: FollowupIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Multi-turn follow-up: creates a new consultation linked to the previous context."""
    try:
        row = _run_async(
            legal_workspace_module.create_consultation_followup(
                db, current_user, consultation_id=item_id, question=req.question
            )
        )
    except LookupError as exc:
        if str(exc) == "LEGAL_CONSULTATION_NOT_FOUND":
            raise api_error(404, "原咨询记录不存在", code="LEGAL_CONSULTATION_NOT_FOUND")
        raise
    except ValueError as exc:
        if str(exc) == "QUOTA_EXCEEDED":
            raise api_error(429, "本月咨询配额已用完，请升级订阅", code="QUOTA_EXCEEDED")
        raise
    return serialize(row)


@router.get("/consultations")
def list_consultations(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return legal_workspace_read_module.list_rows(db, current_user, "consultation")


@router.get("/consultations/{item_id}")
def get_consultation(item_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        return serialize(legal_workspace_read_module.get_row(db, current_user, "consultation", item_id))
    except LookupError:
        raise api_error(404, "法律咨询记录不存在", code="LEGAL_CONSULTATION_NOT_FOUND")


@router.post("/contract-reviews")
@router.post("/contracts/analyze")
def create_contract_review(req: ContractReviewIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        row = _run_async(
            legal_workspace_module.create_contract_review(
                db, current_user, title=req.title, content=req.content,
                document_id=req.document_id, review_policy_id=req.review_policy_id,
                review_policy_override=req.review_policy_override, case_id=req.case_id,
            )
        )
    except ValueError as exc:
        if str(exc) == "QUOTA_EXCEEDED":
            raise api_error(429, "本月合同审查配额已用完，请升级订阅", code="QUOTA_EXCEEDED")
        raise
    except LookupError as exc:
        if str(exc) == "LEGAL_REVIEW_POLICY_NOT_FOUND":
            raise api_error(404, "审查策略不存在或不可用", code="LEGAL_REVIEW_POLICY_NOT_FOUND")
        if str(exc) == "LEGAL_CASE_NOT_FOUND":
            raise api_error(404, "案件不存在或无权访问", code="LEGAL_CASE_NOT_FOUND")
        raise
    return serialize(row)


@router.get("/contract-reviews")
def list_contract_reviews(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return legal_workspace_read_module.list_rows(db, current_user, "contract_review")


@router.get("/contract-reviews/{item_id}")
def get_contract_review(item_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        return serialize(legal_workspace_read_module.get_row(db, current_user, "contract_review", item_id))
    except LookupError:
        raise api_error(404, "合同审查记录不存在", code="LEGAL_CONTRACT_REVIEW_NOT_FOUND")


@router.get("/contract-reviews/{item_id}/versions")
def list_contract_review_versions(item_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        return legal_workspace_read_module.versions(db, current_user, "contract_review", item_id)
    except LookupError:
        raise api_error(404, "合同审查记录不存在", code="LEGAL_CONTRACT_REVIEW_NOT_FOUND")


@router.post("/contract-reviews/{item_id}/resubmit")
def resubmit_contract_review(item_id: int, req: ContractReviewIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """针对被退回的合同审查修改内容后重新提交，旧版本自动存档（FL.md 6.2 版本留痕）。"""
    try:
        row = _run_async(
            legal_workspace_module.resubmit_contract_review(
                db, current_user, review_id=item_id, title=req.title, content=req.content
            )
        )
    except LookupError as exc:
        if str(exc) == "LEGAL_CONTRACT_REVIEW_NOT_FOUND":
            raise api_error(404, "合同审查记录不存在", code="LEGAL_CONTRACT_REVIEW_NOT_FOUND")
        raise
    except ValueError as exc:
        if str(exc) == "LEGAL_CONTRACT_REVIEW_RESUBMIT_INVALID_STATUS":
            raise api_error(400, "仅退回补充事实的记录可重新提交", code="LEGAL_CONTRACT_REVIEW_RESUBMIT_INVALID_STATUS")
        raise
    return serialize(row)


class ContractCompareIn(BaseModel):
    title_a: str = Field(default="合同A", max_length=256)
    content_a: str = Field(min_length=1, max_length=50000)
    title_b: str = Field(default="合同B", max_length=256)
    content_b: str = Field(min_length=1, max_length=50000)


@router.post("/contract-compare")
def compare_contracts_endpoint(req: ContractCompareIn, current_user: User = Depends(get_current_user)):
    result = _run_async(
        compare_contracts(
            req.content_a, req.content_b,
            title_a=req.title_a, title_b=req.title_b,
            user_id=current_user.id,
        )
    )
    return result


@router.post("/contract-reviews/upload")
def upload_contract_review(
    file: UploadFile = File(...),
    title: str | None = None,
    case_id: int | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload a contract file (PDF/DOCX/TXT) and run review."""
    allowed_exts = {".pdf", ".docx", ".doc", ".txt", ".md"}
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in allowed_exts:
        raise api_error(400, f"不支持的文件类型: {ext}，支持 PDF/DOCX/TXT/MD", code="LEGAL_CONTRACT_FILE_TYPE_INVALID")

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            content = file.file.read()
            tmp.write(content)
            tmp_path = tmp.name

        from app.services.document_service import extract_file_text
        file_type = ext.lstrip(".")
        text = extract_file_text(tmp_path, file_type)
        if not text or len(text.strip()) < 20:
            raise api_error(400, "文件内容为空或解析失败，请直接粘贴合同文本", code="LEGAL_CONTRACT_FILE_PARSE_FAILED")

        contract_title = title or os.path.splitext(file.filename or "上传合同")[0]
        try:
            row = _run_async(
                legal_workspace_module.create_contract_review(
                    db, current_user, title=contract_title, content=text, case_id=case_id,
                )
            )
        except ValueError as exc:
            if str(exc) == "QUOTA_EXCEEDED":
                raise api_error(429, "本月合同审查配额已用完，请升级订阅", code="QUOTA_EXCEEDED")
            raise
        except LookupError as exc:
            if str(exc) == "LEGAL_CASE_NOT_FOUND":
                raise api_error(404, "案件不存在或无权访问", code="LEGAL_CASE_NOT_FOUND")
            raise
        return serialize(row)
    except HTTPException:
        raise
    except Exception as exc:
        if should_passthrough_exception(exc):
            raise
        raise api_error(500, "合同文件审查失败", code="LEGAL_CONTRACT_UPLOAD_FAILED", detail=str(exc))
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@router.get("/document-templates")
def document_templates():
    return legal_workspace_read_module.templates()


@router.post("/drafts")
def create_draft(req: DraftIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        row, missing_required = _run_async(
            legal_workspace_module.create_draft(
                db, current_user, document_type=req.document_type, fields=req.fields, case_id=req.case_id
            )
        )
    except ValueError as exc:
        if str(exc) == "QUOTA_EXCEEDED":
            raise api_error(429, "本月文书生成配额已用完，请升级订阅", code="QUOTA_EXCEEDED")
        raise
    except KeyError as exc:
        if exc.args and exc.args[0] == "LEGAL_DRAFT_TYPE_INVALID":
            raise api_error(400, "暂不支持该文书类型", code="LEGAL_DRAFT_TYPE_INVALID")
        raise
    except LookupError as exc:
        if str(exc) == "LEGAL_CASE_NOT_FOUND":
            raise api_error(404, "案件不存在或无权访问", code="LEGAL_CASE_NOT_FOUND")
        raise
    result = serialize(row)
    if missing_required:
        result["missing_required"] = missing_required
    return result


@router.get("/drafts")
def list_drafts(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return legal_workspace_read_module.list_rows(db, current_user, "draft")


@router.get("/drafts/{item_id}")
def get_draft(item_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        return serialize(legal_workspace_read_module.get_row(db, current_user, "draft", item_id))
    except LookupError:
        raise api_error(404, "法律文书草稿不存在", code="LEGAL_DRAFT_NOT_FOUND")


@router.get("/drafts/{item_id}/export/docx")
def export_legal_draft_docx(item_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """V3.1：将当前用户的法律文书草稿导出为 DOCX，不暴露服务器文件路径。"""
    row = db.query(LegalDraft).filter(LegalDraft.id == item_id, LegalDraft.user_id == current_user.id).first()
    if not row:
        raise api_error(404, "法律文书不存在", code="LEGAL_DRAFT_NOT_FOUND")
    from docx import Document as WordDocument
    output_dir = Path(get_settings().STORAGE_LOCAL_DIR) / "generated" / "legal-drafts"
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"legal_draft_{row.id}_v{row.version}.docx"
    path = output_dir / filename
    document = WordDocument()
    document.add_heading(row.title, level=0)
    for paragraph in (row.content or "").splitlines():
        document.add_paragraph(paragraph)
    document.save(path)
    audit.log(db, current_user, "legal_draft_export_docx", target_type="draft", target_id=row.id, detail=f"version={row.version}")
    return FileResponse(path, filename=filename,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.get("/drafts/{item_id}/versions")
def list_draft_versions(item_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        return legal_workspace_read_module.versions(db, current_user, "draft", item_id)
    except LookupError:
        raise api_error(404, "法律文书草稿不存在", code="LEGAL_DRAFT_NOT_FOUND")


@router.post("/drafts/{item_id}/resubmit")
def resubmit_draft(item_id: int, req: DraftIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """针对被退回补充事实的文书草稿修改字段后重新提交，旧版本自动存档。"""
    try:
        row, missing_required = _run_async(
            legal_workspace_module.resubmit_draft(
                db, current_user, draft_id=item_id, document_type=req.document_type, fields=req.fields
            )
        )
    except LookupError as exc:
        if str(exc) == "LEGAL_DRAFT_NOT_FOUND":
            raise api_error(404, "法律文书草稿不存在", code="LEGAL_DRAFT_NOT_FOUND")
        raise
    except ValueError as exc:
        if str(exc) == "LEGAL_DRAFT_RESUBMIT_INVALID_STATUS":
            raise api_error(400, "仅待补充事实或退回补充的草稿可重新提交", code="LEGAL_DRAFT_RESUBMIT_INVALID_STATUS")
        raise
    except KeyError as exc:
        if exc.args and exc.args[0] == "LEGAL_DRAFT_TYPE_INVALID":
            raise api_error(400, "暂不支持该文书类型", code="LEGAL_DRAFT_TYPE_INVALID")
        raise
    result = serialize(row)
    if missing_required:
        result["missing_required"] = missing_required
    return result


@router.get("/review-queue")
def review_queue(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return legal_workspace_read_module.review_queue(db, current_user)


@router.post("/review-queue/{target_type}/{target_id}/actions")
def review_action(target_type: str, target_id: int, req: ReviewActionIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        return legal_workspace_read_module.apply_review_action(
            db, current_user, target_type=target_type, target_id=target_id,
            action=req.action, note=req.note,
        )
    except LookupError:
        raise api_error(404, "待审核记录不存在", code="LEGAL_REVIEW_TARGET_NOT_FOUND")
    except PermissionError:
        raise api_error(403, "仅审核律师或管理员可执行审核动作", code="LEGAL_REVIEW_FORBIDDEN")
    except ValueError:
        raise api_error(400, "不支持的审核动作", code="LEGAL_REVIEW_ACTION_INVALID")


@router.post("/review-queue/{target_type}/{target_id}/comments")
def add_review_comment(target_type: str, target_id: int, req: ReviewCommentIn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """独立批注：不改变审核状态，仅留言，供提交人和审核律师随时沟通（FL.md 6.2 律师批注）。"""
    try:
        return legal_workspace_read_module.add_review_comment(
            db, current_user, target_type=target_type, target_id=target_id, note=req.note,
        )
    except LookupError:
        raise api_error(404, "记录不存在", code="LEGAL_REVIEW_TARGET_NOT_FOUND")
    except PermissionError:
        raise api_error(403, "仅提交人或审核律师可添加批注", code="LEGAL_REVIEW_COMMENT_FORBIDDEN")


@router.get("/review-queue/{target_type}/{target_id}/history")
def review_history(target_type: str, target_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """查看单条记录的审核历史（并排展示原文与历次审核意见）。"""
    try:
        return legal_workspace_read_module.review_history(
            db, current_user, target_type=target_type, target_id=target_id,
        )
    except LookupError:
        raise api_error(404, "记录不存在", code="LEGAL_REVIEW_TARGET_NOT_FOUND")
    except PermissionError:
        raise api_error(403, "无权查看该记录的审核历史", code="LEGAL_REVIEW_HISTORY_FORBIDDEN")


@router.get("/review-stats")
def review_stats(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """审核历史统计：退回原因分布、审核动作分布、平均处理时长（FL.md 6.5）。"""
    try:
        return legal_workspace_read_module.review_stats(db, current_user)
    except PermissionError:
        raise api_error(403, "仅审核律师或管理员可查看审核统计", code="LEGAL_REVIEW_STATS_FORBIDDEN")


# ── 团队共享知识库 API（Phase 9）─────────────────────────────────────────────

from app.models.org import LegalMemberRole, OrganizationMember  # noqa: E402


def _require_org_member_legal(db: Session, user_id: int, org_id: int) -> OrganizationMember:
    member = db.query(OrganizationMember).filter(
        OrganizationMember.organization_id == org_id,
        OrganizationMember.user_id == user_id,
    ).first()
    if not member:
        raise api_error(403, "不是该组织成员", code="NOT_ORG_MEMBER")
    return member


def _require_editor_legal(db: Session, user_id: int, org_id: int) -> OrganizationMember:
    member = _require_org_member_legal(db, user_id, org_id)
    if member.legal_role == LegalMemberRole.client.value:
        raise api_error(403, "客户角色无法管理团队知识库", code="INSUFFICIENT_ROLE")
    return member


@router.get("/orgs/{org_id}/shared-sources")
def list_team_sources(
    org_id: int,
    source_type: str | None = None,
    law_area: str | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """列出组织团队共享法源库（所有成员可读）。"""
    _require_org_member_legal(db, current_user.id, org_id)
    q = db.query(LegalSource).filter(
        LegalSource.organization_id == org_id,
        LegalSource.scope == "team",
    )
    if source_type:
        q = q.filter(LegalSource.source_type == source_type)
    sources = q.order_by(LegalSource.updated_at.desc()).all()
    if law_area:
        sources = [
            s for s in sources
            if law_area in (json.loads(s.law_area_json) if s.law_area_json else [])
        ]
    return [_serialize_source(s) for s in sources]


@router.post("/orgs/{org_id}/shared-sources/{source_id}/share", status_code=200)
def share_source_to_team(
    org_id: int,
    source_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """将个人法源共享给团队（editor/reviewer/admin 可操作）。"""
    _require_editor_legal(db, current_user.id, org_id)
    source = db.query(LegalSource).filter(
        LegalSource.id == source_id,
        LegalSource.user_id == current_user.id,
    ).first()
    if not source:
        raise api_error(404, "法源不存在或不属于当前用户", code="LEGAL_SOURCE_NOT_FOUND")
    source.organization_id = org_id
    source.scope = "team"
    db.commit()
    db.refresh(source)
    return _serialize_source(source)


@router.delete("/orgs/{org_id}/shared-sources/{source_id}/share", status_code=200)
def unshare_source_from_team(
    org_id: int,
    source_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """取消团队共享，将法源改回个人（原始创建者或管理员可操作）。"""
    member = _require_org_member_legal(db, current_user.id, org_id)
    source = db.query(LegalSource).filter(
        LegalSource.id == source_id,
        LegalSource.organization_id == org_id,
        LegalSource.scope == "team",
    ).first()
    if not source:
        raise api_error(404, "未找到该组织下的共享法源", code="LEGAL_SOURCE_NOT_FOUND")
    is_owner = source.user_id == current_user.id
    is_admin = member.legal_role == LegalMemberRole.admin.value
    if not is_owner and not is_admin:
        raise api_error(403, "只有法源创建者或组织管理员可取消共享", code="INSUFFICIENT_ROLE")
    source.organization_id = None
    source.scope = "personal"
    db.commit()
    db.refresh(source)
    return _serialize_source(source)


# ── 用户反馈闭环 ──────────────────────────────────────────────────────────────


class FeedbackIn(BaseModel):
    score: int = Field(..., description="满意度评分：1=满意，-1=不满意")
    note: str | None = Field(default=None, max_length=2000)


@router.post("/consultations/{item_id}/feedback")
def consultation_feedback(
    item_id: int,
    req: FeedbackIn,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """提交法律咨询满意度反馈（👍/👎）。"""
    if req.score not in (1, -1):
        raise api_error(400, "score 只能是 1 或 -1", code="FEEDBACK_SCORE_INVALID")
    row = db.query(LegalConsultation).filter(
        LegalConsultation.id == item_id,
        LegalConsultation.user_id == current_user.id,
    ).first()
    if not row:
        raise api_error(404, "法律咨询记录不存在", code="LEGAL_CONSULTATION_NOT_FOUND")
    row.feedback_score = req.score
    row.feedback_note = req.note
    db.commit()
    return {"id": row.id, "feedback_score": row.feedback_score}


@router.post("/contract-reviews/{item_id}/feedback")
def contract_review_feedback(
    item_id: int,
    req: FeedbackIn,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """提交合同审查满意度反馈（👍/👎）。"""
    if req.score not in (1, -1):
        raise api_error(400, "score 只能是 1 或 -1", code="FEEDBACK_SCORE_INVALID")
    row = db.query(ContractReview).filter(
        ContractReview.id == item_id,
        ContractReview.user_id == current_user.id,
    ).first()
    if not row:
        raise api_error(404, "合同审查记录不存在", code="LEGAL_CONTRACT_REVIEW_NOT_FOUND")
    row.feedback_score = req.score
    row.feedback_note = req.note
    db.commit()
    return {"id": row.id, "feedback_score": row.feedback_score}


@router.post("/drafts/{item_id}/feedback")
def draft_feedback(
    item_id: int,
    req: FeedbackIn,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """提交文书草稿满意度反馈（👍/👎）。"""
    if req.score not in (1, -1):
        raise api_error(400, "score 只能是 1 或 -1", code="FEEDBACK_SCORE_INVALID")
    row = db.query(LegalDraft).filter(
        LegalDraft.id == item_id,
        LegalDraft.user_id == current_user.id,
    ).first()
    if not row:
        raise api_error(404, "法律文书草稿不存在", code="LEGAL_DRAFT_NOT_FOUND")
    row.feedback_score = req.score
    row.feedback_note = req.note
    db.commit()
    return {"id": row.id, "feedback_score": row.feedback_score}
